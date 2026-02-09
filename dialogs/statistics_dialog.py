import calendar
from PySide6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel, 
                             QTableWidget, QTableWidgetItem, QHeaderView, 
                             QPushButton, QAbstractItemView, QFileDialog, 
                             QMessageBox, QGroupBox, QDateEdit)
from PySide6.QtCore import Qt, QDate
from PySide6.QtGui import QGuiApplication, QColor


class StatisticsDialog(QDialog):
    def __init__(self, parent, db_conn):
        super().__init__(parent)
        self.setWindowTitle("Báo cáo thống kê đào tạo")
        self.resize(1000, 750)
        self.db_conn = db_conn
        
        self.setStyleSheet("""
            QDialog { background-color: #ffffff; }
            QLabel#Title { font-size: 18px; font-weight: bold; color: #2c3e50; }
            QTableWidget { border: 1px solid #dee2e6; gridline-color: #eee; }
            QHeaderView::section { background-color: #f8f9fa; font-weight: bold; border: 1px solid #dee2e6; }
            QPushButton#ExportBtn { background-color: #2b5797; color: white; font-weight: bold; }
        """)

        layout = QVBoxLayout(self)
        
        title = QLabel("BÁO CÁO THỐNG KÊ TÌNH HÌNH LỚP HỌC")
        title.setObjectName("Title")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        filter_group = QGroupBox("Chọn khoảng thời gian thống kê")
        filter_layout = QHBoxLayout(filter_group)
        
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDate(QDate.currentDate().addDays(-7))# Ngày bắt đầu lọc mặc định là 7 ngày trước ngày hiện tại
        
        self.end_date = QDateEdit()
        self.end_date.setCalendarPopup(True)
        self.end_date.setDate(QDate.currentDate())
        
        btn_refresh = QPushButton("🔄 Tải dữ liệu")
        btn_refresh.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold; font-size :14px;")
        btn_refresh.clicked.connect(self.calculate_stats)
        
        filter_layout.addWidget(QLabel("Từ ngày:"))
        filter_layout.addWidget(self.start_date)
        filter_layout.addWidget(QLabel("Đến ngày:"))
        filter_layout.addWidget(self.end_date)
        filter_layout.addWidget(btn_refresh)
        filter_layout.addStretch()
        layout.addWidget(filter_group)

        # Bảng 1: Thống kê tóm tắt
        layout.addWidget(QLabel("<b>1. Tóm tắt tỉ lệ chuyên cần theo lớp:</b> (Di chuột bôi đen và nhấn Ctrl+C để copy)"))
        self.summary_table = QTableWidget(0, 5)
        self.summary_table.setHorizontalHeaderLabels(["Tên lớp", "Sĩ số", "Số bạn đi học", "Số bạn nghỉ", "Tỉ lệ (%)"])
        
        self.summary_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.summary_table.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.summary_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.summary_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.summary_table.setFixedHeight(200)
        layout.addWidget(self.summary_table)

        # Bảng 2: Thống kê số buổi học theo ID học sinh
        layout.addWidget(QLabel("<b>2. Thống kê số buổi học theo ID học sinh:</b>"))
        self.student_stats_table = QTableWidget(0, 6)
        self.student_stats_table.setHorizontalHeaderLabels([
            "ID học sinh", "Họ tên", "Lớp", "Số buổi đi học", "Số buổi nghỉ", "Tổng số buổi"
        ])
        self.student_stats_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.student_stats_table.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.student_stats_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.student_stats_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.student_stats_table.setFixedHeight(250)
        layout.addWidget(self.student_stats_table)

        # Bảng 3: Chi tiết học sinh & Nhận xét
        layout.addWidget(QLabel("<b>3. Chi tiết học viên và nhận xét:</b>"))
        self.detail_table = QTableWidget()
        self.detail_table.setColumnCount(3)
        self.detail_table.setHorizontalHeaderLabels([
            "Họ tên", "Lớp đang học", "Nhận xét"
        ])
        self.detail_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.detail_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.detail_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        layout.addWidget(self.detail_table)
        
        btn_box = QHBoxLayout()
        self.btn_export_word = QPushButton("📄 XUẤT FILE WORD")
        self.btn_export_word.setObjectName("ExportBtn")
        self.btn_export_word.setFixedSize(180, 40)
        self.btn_export_word.clicked.connect(self.export_to_word)
        
        btn_close = QPushButton("Đóng")
        btn_close.setFixedSize(100, 40)
        btn_close.clicked.connect(self.close)
        
        btn_box.addStretch()
        btn_box.addWidget(self.btn_export_word)
        btn_box.addWidget(btn_close)
        layout.addLayout(btn_box)

        self.calculate_stats()
    
    def keyPressEvent(self, event):
        #Xử lý bôi đen rồi Ctrl+C để copy từ bảng (cho báo cáo)
        if event.modifiers() == Qt.ControlModifier and event.key() == Qt.Key_C:
            table = self.focusWidget()
            if isinstance(table, QTableWidget):
                selected_ranges = table.selectedRanges()
                if not selected_ranges:
                    return

                text = ""
                for r_range in selected_ranges:
                    for r in range(r_range.topRow(), r_range.bottomRow() + 1):
                        row_data = []
                        for c in range(r_range.leftColumn(), r_range.rightColumn() + 1):
                            item = table.item(r, c)
                            row_data.append(item.text() if item else "")
                        text += "\t".join(row_data) + "\n"
                
                QGuiApplication.clipboard().setText(text)
        else:
            super().keyPressEvent(event)

    def count_weekends_in_month(self, year, month):
        count = 0
        num_days = calendar.monthrange(year, month)[1]
        for day in range(1, num_days + 1):
            # 5 là Thứ 7, 6 là Chủ nhật trong thư viện calendar
            if calendar.weekday(year, month, day) in [5, 6]:
                count += 1
        return count
    
    def calculate_stats(self):
        d1 = self.start_date.date().toString("yyyy-MM-dd")
        d2 = self.end_date.date().toString("yyyy-MM-dd")
        cursor = self.db_conn.cursor()
        
        classes = ["Sáng T7", "Chiều T7", "Sáng CN", "Chiều CN"]
        
        # 1. Bảng tóm tắt
        self.summary_table.setRowCount(0)
        from database import Database
        db_temp = Database()
        
        for class_name in classes:
            # Sĩ số: Lấy số lượng học sinh của từng lớp đang có
            cursor.execute("SELECT COUNT(DISTINCT name) FROM progress WHERE class_name = ? AND date BETWEEN ? AND ?", 
                          (class_name, d1, d2))
            total_students = cursor.fetchone()[0] or 0
            
            # Tính số buổi cho lớp này trong khoảng thời gian đã chọn
            expected_sessions = db_temp.count_expected_sessions(class_name, d1, d2)
            
            # Tổng số buổi= sĩ số × số buổi
            total_expected = total_students * expected_sessions
            
            # Số buổi đi học thực tế
            cursor.execute("SELECT COUNT(*) FROM progress WHERE class_name = ? AND status = 'Đi học' AND date BETWEEN ? AND ?", 
                           (class_name, d1, d2))
            present = cursor.fetchone()[0] or 0
            
            # Số buổi nghỉ = tổng dự kiến - số buổi đi
            absent = max(0, total_expected - present)
            
            # Tỉ lệ % = (số buổi đi / tổng dự kiến) × 100
            percent = (present / total_expected * 100) if total_expected > 0 else 0

            row = self.summary_table.rowCount()
            self.summary_table.insertRow(row)
            self.summary_table.setItem(row, 0, QTableWidgetItem(class_name))
            self.summary_table.setItem(row, 1, QTableWidgetItem(str(total_students)))
            self.summary_table.setItem(row, 2, QTableWidgetItem(str(present)))
            self.summary_table.setItem(row, 3, QTableWidgetItem(str(absent)))
            self.summary_table.setItem(row, 4, QTableWidgetItem(f"{percent:.1f}%"))

        # 2. Bảng thống kê số buổi học theo ID học sinh
        self.student_stats_table.setRowCount(0)
        from database import Database
        db = Database()
        students_with_attendance = db.get_all_students_with_attendance(d1, d2)
        
        for student_id, name, class_name, attended, absent, total in students_with_attendance:
            row = self.student_stats_table.rowCount()
            self.student_stats_table.insertRow(row)
            self.student_stats_table.setItem(row, 0, QTableWidgetItem(student_id))
            self.student_stats_table.setItem(row, 1, QTableWidgetItem(name))
            self.student_stats_table.setItem(row, 2, QTableWidgetItem(class_name))
            
            # Tô màu cho số buổi đi học
            attended_item = QTableWidgetItem(str(attended))
            attended_item.setForeground(QColor("#28a745"))
            attended_item.setTextAlignment(Qt.AlignCenter)
            self.student_stats_table.setItem(row, 3, attended_item)
            
            # Tô màu cho số buổi nghỉ
            absent_item = QTableWidgetItem(str(absent))
            absent_item.setForeground(QColor("#dc3545"))
            absent_item.setTextAlignment(Qt.AlignCenter)
            self.student_stats_table.setItem(row, 4, absent_item)
            
            # Tổng số buổi
            total_item = QTableWidgetItem(str(total))
            total_item.setTextAlignment(Qt.AlignCenter)
            self.student_stats_table.setItem(row, 5, total_item)

        # 3. Bảng chi tiết
        self.detail_table.setRowCount(0)
        
        # Lấy danh sách học sinh từ bảng students với các lớp của họ
        cursor.execute("""
            SELECT DISTINCT s.name, s.class_name
            FROM students s
            ORDER BY s.name
        """)
        
        students_list = cursor.fetchall()
        
        for s_name, s_class in students_list:
            # Lấy tất cả các lớp học sinh tham gia trong khoảng thời gian
            cursor.execute("""
                SELECT DISTINCT class_name 
                FROM progress 
                WHERE name = ? AND date BETWEEN ? AND ?
                ORDER BY class_name
            """, (s_name, d1, d2))
            
            classes_in_period = [row[0] for row in cursor.fetchall()]
            display_classes = ", ".join(classes_in_period) if classes_in_period else s_class

            row = self.detail_table.rowCount()
            self.detail_table.insertRow(row)
            self.detail_table.setItem(row, 0, QTableWidgetItem(s_name))
            self.detail_table.setItem(row, 1, QTableWidgetItem(display_classes))
            self.detail_table.setItem(row, 2, QTableWidgetItem(""))  # Nhận xét để trống

    def export_to_word(self):
        #Xuất báo cáo ra file Word
        from docx import Document
        
        path, _ = QFileDialog.getSaveFileName(
            self, 
            "Lưu báo cáo", 
            f"Bao_cao_hoc_tap_{QDate.currentDate().toString('ddMMyy')}.docx", 
            "Word Files (*.docx)"
        )
        if not path:
            return

        try:
            doc = Document()
            doc.add_heading('BÁO CÁO TÌNH HÌNH HỌC TẬP', 0)
            
            # Thông tin thời gian
            time_str = f"Từ ngày: {self.start_date.date().toString('dd/MM/yyyy')} đến ngày: {self.end_date.date().toString('dd/MM/yyyy')}"
            doc.add_paragraph(time_str)

            # 1. Bảng tóm tắt
            doc.add_heading('1. Thống kê chuyên cần theo lớp', level=1)
            table1 = doc.add_table(rows=1, cols=5)
            table1.style = 'Table Grid'
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = 'Tên lớp'
            hdr_cells[1].text = 'Sĩ số'
            hdr_cells[2].text = 'Đi học'
            hdr_cells[3].text = 'Nghỉ học'
            hdr_cells[4].text = 'Tỉ lệ (%)'

            for r in range(self.summary_table.rowCount()):
                row_cells = table1.add_row().cells
                for c in range(5):
                    row_cells[c].text = self.summary_table.item(r, c).text()

            doc.add_paragraph("\n")

            # 2. Bảng thống kê số buổi học theo ID
            doc.add_heading('2. Thống kê số buổi học theo ID học sinh', level=1)
            table_students = doc.add_table(rows=1, cols=6)
            table_students.style = 'Table Grid'
            hdr_students = table_students.rows[0].cells
            hdr_students[0].text = 'ID học sinh'
            hdr_students[1].text = 'Họ tên'
            hdr_students[2].text = 'Lớp'
            hdr_students[3].text = 'Số buổi đi học'
            hdr_students[4].text = 'Số buổi nghỉ'
            hdr_students[5].text = 'Tổng số buổi'

            for r in range(self.student_stats_table.rowCount()):
                row_cells = table_students.add_row().cells
                for c in range(6):
                    row_cells[c].text = self.student_stats_table.item(r, c).text()

            doc.add_paragraph("\n")

            # 3. Bảng chi tiết
            doc.add_heading('3. Chi tiết học viên và nhận xét', level=1)
            table2 = doc.add_table(rows=1, cols=3)
            table2.style = 'Table Grid'
            hdr_cells2 = table2.rows[0].cells
            hdr_cells2[0].text = 'Họ tên'
            hdr_cells2[1].text = 'Lớp đang học'
            hdr_cells2[2].text = 'Nhận xét'

            for r in range(self.detail_table.rowCount()):
                row_cells = table2.add_row().cells
                for c in range(3):
                    item = self.detail_table.item(r, c)
                    row_cells[c].text = item.text() if item else ""

            doc.save(path)
            QMessageBox.information(self, "Thành công", f"Đã xuất báo cáo tại:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Không thể xuất file: {str(e)}")
