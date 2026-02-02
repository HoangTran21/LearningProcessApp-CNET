import sys
import sqlite3
from datetime import datetime
from PySide6.QtWidgets import (QApplication, QMainWindow, QTableWidget, 
                             QTableWidgetItem, QVBoxLayout, QHBoxLayout, 
                             QWidget, QPushButton, QHeaderView, QComboBox, 
                             QLineEdit, QTextEdit, QLabel, QDialog, QFormLayout, 
                             QMessageBox, QGroupBox, QCheckBox, QDateEdit,
                             QAbstractItemView, QCompleter, QListWidget, QListWidgetItem, QFileDialog)
from PySide6.QtCore import Qt, QDate
from PySide6.QtGui import QColor, QIcon, QScreen, QGuiApplication
from docx import Document

class EntryDialog(QDialog):
    def __init__(self, parent=None, data=None, student_list=None, db_conn=None):
        super().__init__(parent)
        self.setWindowTitle("Thông tin tiến độ")
        self.setFixedWidth(500)
        self.setWindowIcon(QIcon("logo_app.png"))
        self.db_conn = db_conn
        
        self.setStyleSheet("""
            QDialog { background-color: #ffffff; }
            QGroupBox { 
                font-weight: bold; border: 1px solid #d1d1d1;
                border-radius: 8px; background-color: #ffffff; 
                margin-top: 25px;
            }
            QLineEdit, QComboBox, QDateEdit, QTextEdit { 
                border: 1px solid #ccc; border-radius: 4px; padding: 8px; background-color: white;
            }
        """)

        layout = QVBoxLayout(self)
        form_group = QGroupBox("Chi tiết học tập")
        form_layout = QFormLayout(form_group)
        form_layout.setContentsMargins(10, 25, 10, 10)

        self.date_edit = QDateEdit()
        self.date_edit.setCalendarPopup(True)
        self.date_edit.setDisplayFormat("yyyy-MM-dd")
        self.date_edit.setDate(QDate.fromString(data[1], "yyyy-MM-dd") if data else QDate.currentDate())

        self.name_edit = QLineEdit(data[2] if data else "")
        self.class_edit = QComboBox()
        self.class_edit.addItems(["Sáng T7", "Chiều T7", "Sáng CN", "Chiều CN"])
        if data: self.class_edit.setCurrentText(data[3])

        if student_list:
            completer = QCompleter(student_list)
            completer.setCaseSensitivity(Qt.CaseInsensitive)
            self.name_edit.setCompleter(completer)
            self.name_edit.editingFinished.connect(lambda: self.auto_fill_class(self.name_edit.text()))
        
        self.status_edit = QComboBox()
        self.status_edit.addItems(["Đi học", "Nghỉ học"])
        if data: self.status_edit.setCurrentText(data[4])
        
        self.content_edit = QTextEdit(data[5] if data else "")
        self.content_edit.setPlaceholderText("Nhập nội dung bài học...")
        self.content_edit.setMinimumHeight(150)

        self.highlight_cb = QComboBox()
        self.highlight_cb.addItems(["Bình thường", "Cần chú ý", "Học tốt", "Báo động"])
        if data: self.highlight_cb.setCurrentIndex(data[6])

        form_layout.addRow("Ngày:", self.date_edit)
        form_layout.addRow("Học sinh:", self.name_edit)
        form_layout.addRow("Lớp:", self.class_edit)
        form_layout.addRow("Trạng thái:", self.status_edit)
        form_layout.addRow("Nội dung:", self.content_edit)
        form_layout.addRow("Đánh giá:", self.highlight_cb)
        
        layout.addWidget(form_group)
        self.btn_save = QPushButton("LƯU DỮ LIỆU")
        self.btn_save.setFixedHeight(40)
        self.btn_save.setStyleSheet("background-color: #007bff; color: white; font-weight: bold; border-radius: 4px;")
        self.btn_save.clicked.connect(self.validate_and_accept)
        layout.addWidget(self.btn_save)

    def auto_fill_class(self, name):
        if self.db_conn and name:
            cursor = self.db_conn.cursor()
            cursor.execute("SELECT class_name FROM progress WHERE name = ? ORDER BY date DESC LIMIT 1", (name,))
            res = cursor.fetchone()
            if res: self.class_edit.setCurrentText(res[0])

    def validate_and_accept(self):
        if not self.name_edit.text().strip():
            QMessageBox.warning(self, "Cảnh báo", "Vui lòng nhập tên học sinh!")
            return
        self.accept()

    def get_data(self):
        return (self.date_edit.date().toString("yyyy-MM-dd"), self.name_edit.text().strip(), 
                self.class_edit.currentText(), self.status_edit.currentText(), 
                self.content_edit.toPlainText(), self.highlight_cb.currentIndex())

class AttendanceDialog(QDialog):
    def __init__(self, parent, db_conn):
        super().__init__(parent)
        self.setWindowTitle("Điểm danh học sinh")
        self.setFixedSize(500, 680)
        self.db_conn = db_conn
        self.all_students = []
        
        self.center_dialog()

        self.setStyleSheet("""
            QDialog { background-color: #f8f9fa; }
            QGroupBox { 
                font-weight: bold; border: 1px solid #d1d1d1; 
                border-radius: 8px; margin-top: 15px; background-color: white;
            }
            QListWidget { border: none; outline: none; background: white; }
            
            QListWidget::item { 
                padding: 12px; 
                border-bottom: 1px solid #f0f0f0; 
                color: #333; 
            }
            QListWidget::item:hover { background-color: #f1faff; }
            
            QCheckBox { font-weight: bold; color: #007bff; }
            
            QListWidget::indicator { width: 20px; height: 20px; border: 2px solid #ccc; border-radius: 4px; }
            QListWidget::indicator:checked { background-color: #28a745; border: 2px solid #28a745; }
            
            QCheckBox#SelectAll {
                font-size: 16px;
                spacing: 8px;
                color: #d63384;
                font-weight: bold;
                padding: 8px;
                background-color: #ffffff;
                border-radius: 5px;
            }
            QCheckBox#SelectAll:hover {
                background-color: #d63384;
                color: #ffffff;
            }
            QCheckBox#SelectAll::indicator {
                width: 20px;
                height: 20px;
                border: 2px solid #d63384;
                border-radius: 4px;
                background-color: white;
            }

            QCheckBox#SelectAll::indicator:hover {
                border: 2px solid #b02a6f;
                background-color: #f8d7da;
            }

            QCheckBox#SelectAll::indicator:checked {
                background-color: #28a745; /* Nền xanh lá khi chọn */
                border: 2px solid #1e7e34; /* Viền xanh đậm khi chọn */
            }
            
            QCheckBox#SelectAll::indicator:checked:pressed {
                background-color: #1e7e34;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(10)
        
        filter_layout = QHBoxLayout()
        filter_label = QLabel("Lọc lớp học:")
        self.class_filter = QComboBox()
        self.class_filter.addItems(["Tất cả lớp", "Sáng T7", "Chiều T7", "Sáng CN", "Chiều CN"])
        self.class_filter.setStyleSheet("""
            QComboBox { 
                padding: 5px; border: 1px solid #ccc; border-radius: 4px; font-weight: bold; color: #d63384; 
            }
        """)
        self.class_filter.currentTextChanged.connect(self.refresh_list)
        
        filter_layout.addWidget(filter_label)
        filter_layout.addWidget(self.class_filter)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)
        
        date_layout = QHBoxLayout()
        date_label = QLabel("Ngày điểm danh:")
        self.attendance_date = QDateEdit()
        self.attendance_date.setCalendarPopup(True)
        self.attendance_date.setDate(QDate.currentDate())
        self.attendance_date.setStyleSheet("""
            QDateEdit { 
                padding: 5px; border: 1px solid #ccc; border-radius: 4px; font-weight: bold; color: #d63384; 
            }
        """)
        date_layout.addWidget(date_label)
        date_layout.addWidget(self.attendance_date)
        date_layout.addStretch()
        layout.addLayout(date_layout)

        self.cb_select_all = QCheckBox("Chọn tất cả học viên")
        self.cb_select_all.setObjectName("SelectAll") 
        self.cb_select_all.setCursor(Qt.PointingHandCursor)
        self.cb_select_all.stateChanged.connect(self.toggle_select_all)
        layout.addWidget(self.cb_select_all)

        self.group_box = QGroupBox("Danh sách học viên")
        group_layout = QVBoxLayout(self.group_box)
        group_layout.setContentsMargins(10, 25, 10, 10)
        self.list_widget = QListWidget()
        group_layout.addWidget(self.list_widget)
        layout.addWidget(self.group_box)

        btn_box = QHBoxLayout()
        self.btn_confirm = QPushButton("XÁC NHẬN ĐIỂM DANH")
        self.btn_confirm.setFixedHeight(40)
        self.btn_confirm.setStyleSheet("background-color: #28a745; color: white; font-weight: bold; border-radius: 5px;")
        self.btn_confirm.clicked.connect(self.accept)
        
        self.btn_close = QPushButton("Đóng")
        self.btn_close.setFixedHeight(40)
        self.btn_close.clicked.connect(self.reject)
        
        btn_box.addWidget(self.btn_close)
        btn_box.addWidget(self.btn_confirm)
        layout.addLayout(btn_box)

        self.load_students()

    def toggle_select_all(self, state):
        new_state = Qt.Checked if state == 2 else Qt.Unchecked
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.data(Qt.UserRole):
                item.setCheckState(new_state)

    def refresh_list(self):
        self.list_widget.clear()
        self.cb_select_all.blockSignals(True)
        self.cb_select_all.setCheckState(Qt.Unchecked)
        self.cb_select_all.blockSignals(False)
        
        filter_text = self.class_filter.currentText()
        for name, cls in self.all_students:
            if filter_text == "Tất cả lớp" or filter_text == cls:
                item = QListWidgetItem(f"{name}  |  Lớp: {cls}")
                # Đảm bảo UserRole lưu đúng tuple (tên, lớp)
                item.setData(Qt.UserRole, (name, cls))
                # QUAN TRỌNG: Phải có Qt.ItemIsUserCheckable mới hiện tích
                item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable | Qt.ItemIsUserCheckable)
                item.setCheckState(Qt.Unchecked)
                self.list_widget.addItem(item)
        
        self.group_box.setTitle(f"Học viên lớp {filter_text} ({self.list_widget.count()})")

    def get_selected_data(self):
        selected = []
        chosen_date = self.attendance_date.date().toString("yyyy-MM-dd")
        
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.checkState() == Qt.Checked:
                data = item.data(Qt.UserRole)
                if data:
                    selected.append((data[0], data[1], chosen_date))
        return selected

    def center_dialog(self):
        qr = self.frameGeometry()
        cp = QScreen.availableGeometry(QApplication.primaryScreen()).center()
        qr.moveCenter(cp)
        self.move(qr.topLeft())

    def load_students(self):
        try:
            cursor = self.db_conn.cursor()
            cursor.execute("""
                SELECT name, class_name FROM progress 
                WHERE id IN (SELECT MAX(id) FROM progress GROUP BY name)
                ORDER BY class_name, name
            """)
            self.all_students = cursor.fetchall()
            self.refresh_list()
        except Exception as e:
            print(f"Lỗi DB: {e}")

class StatisticsDialog(QDialog):
    def __init__(self, parent, db_conn):
        super().__init__(parent)
        self.setWindowTitle("Báo cáo thống kê đào tạo")
        self.resize(1000, 750)
        self.db_conn = db_conn
        
        self.setStyleSheet("""
            QDialog { background-color: #ffffff; }
            QLabel#Title { font-size: 18px; font-weight: bold; color: #2c3e50; }
            QTableWidget { border: 1px solid #dee2e6; gridline-color: #eee; }
            QHeaderView::section { background-color: #f8f9fa; font-weight: bold; border: 1px solid #dee2e6; }
            QPushButton#ExportBtn { background-color: #2b5797; color: white; font-weight: bold; }
        """)

        layout = QVBoxLayout(self)
        
        # --- Header ---
        title = QLabel("BÁO CÁO THỐNG KÊ TÌNH HÌNH LỚP HỌC")
        title.setObjectName("Title")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        filter_group = QGroupBox("Chọn khoảng thời gian thống kê")
        filter_layout = QHBoxLayout(filter_group)
        
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDate(QDate.currentDate().addDays(-7)) # Mặc định lùi 1 tuần
        
        self.end_date = QDateEdit()
        self.end_date.setCalendarPopup(True)
        self.end_date.setDate(QDate.currentDate())
        
        btn_refresh = QPushButton("🔄 Tải dữ liệu")
        btn_refresh.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold; font-size :14px;")
        btn_refresh.clicked.connect(self.calculate_stats)
        
        filter_layout.addWidget(QLabel("Từ ngày:"))
        filter_layout.addWidget(self.start_date)
        filter_layout.addWidget(QLabel("Đến ngày:"))
        filter_layout.addWidget(self.end_date)
        filter_layout.addWidget(btn_refresh)
        filter_layout.addStretch()
        layout.addWidget(filter_group)

        # --- Bảng 1: Thống kê tóm tắt  
        layout.addWidget(QLabel("<b>1. Tóm tắt tỉ lệ chuyên cần theo lớp:</b> (Di chuột bôi đen và nhấn Ctrl+C để copy)"))
        self.summary_table = QTableWidget(0, 5)
        self.summary_table.setHorizontalHeaderLabels(["Tên lớp", "Sĩ số", "Số bạn đi học", "Số bạn nghỉ", "Tỉ lệ (%)"])
        
        self.summary_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.summary_table.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.summary_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.summary_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.summary_table.setFixedHeight(200)
        layout.addWidget(self.summary_table)

        # --- Bảng 2: Chi tiết học sinh & Nhận xét ---
        layout.addWidget(QLabel("<b>2. Chi tiết học viên và nhận xét:</b>"))
        self.detail_table = QTableWidget()
        self.detail_table.setColumnCount(5) # Tăng lên 5 cột
        self.detail_table.setHorizontalHeaderLabels([
            "Ngày", "Lớp", "Tên học sinh", "Chuyên cần tháng", "Nhận xét cuối buổi"
        ])
        # Điều chỉnh độ rộng cột
        self.detail_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.detail_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.detail_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.detail_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Stretch)
        layout.addWidget(self.detail_table)
        btn_box = QHBoxLayout()
        self.btn_export_word = QPushButton("📄 XUẤT FILE WORD")
        self.btn_export_word.setObjectName("ExportBtn")
        self.btn_export_word.setFixedSize(180, 40)
        self.btn_export_word.clicked.connect(self.export_to_word)
        
        btn_close = QPushButton("Đóng")
        btn_close.setFixedSize(100, 40)
        btn_close.clicked.connect(self.close)
        
        btn_box.addStretch()
        btn_box.addWidget(self.btn_export_word)
        btn_box.addWidget(btn_close)
        layout.addLayout(btn_box)

        self.calculate_stats()
    
    def keyPressEvent(self, event):
        """Xử lý phím Ctrl + C để sao chép vùng bôi đen"""
        if event.modifiers() == Qt.ControlModifier and event.key() == Qt.Key_C:
            table = self.focusWidget()
            if isinstance(table, QTableWidget):
                selected_ranges = table.selectedRanges()
                if not selected_ranges:
                    return

                text = ""
                for r_range in selected_ranges:
                    for r in range(r_range.topRow(), r_range.bottomRow() + 1):
                        row_data = []
                        for c in range(r_range.leftColumn(), r_range.rightColumn() + 1):
                            item = table.item(r, c)
                            row_data.append(item.text() if item else "")
                        text += "\t".join(row_data) + "\n"
                
                QGuiApplication.clipboard().setText(text)
        else:
            super().keyPressEvent(event)

    def count_weekends_in_month(self, year, month):
        """Đếm tổng số ngày Thứ 7 và Chủ nhật trong một tháng cụ thể"""
        import calendar
        count = 0
        num_days = calendar.monthrange(year, month)[1]
        for day in range(1, num_days + 1):
            # 5 là Thứ 7, 6 là Chủ nhật trong thư viện calendar
            if calendar.weekday(year, month, day) in [5, 6]:
                count += 1
        return count
    
    def calculate_stats(self):
        d1 = self.start_date.date().toString("yyyy-MM-dd")
        d2 = self.end_date.date().toString("yyyy-MM-dd")
        cursor = self.db_conn.cursor()
        
        classes = ["Sáng T7", "Chiều T7", "Sáng CN", "Chiều CN"]
        
        # 1. Xử lý Bảng Tóm Tắt
        self.summary_table.setRowCount(0)
        for class_name in classes:
            # Sĩ số: Lấy số lượng học sinh duy nhất của lớp đó trong khoảng thời gian này
            cursor.execute("SELECT COUNT(DISTINCT name) FROM progress WHERE class_name = ?", (class_name,))
            total = cursor.fetchone()[0] or 0
            
            # Số buổi đi học thực tế (tổng các lượt điểm danh 'Đi học')
            cursor.execute("SELECT COUNT(*) FROM progress WHERE class_name = ? AND status = 'Đi học' AND date BETWEEN ? AND ?", 
                           (class_name, d1, d2))
            present = cursor.fetchone()[0] or 0
            
            absent = total - present if total > present else 0
            percent = (present / total * 100) if total > 0 else 0

            row = self.summary_table.rowCount()
            self.summary_table.insertRow(row)
            self.summary_table.setItem(row, 0, QTableWidgetItem(class_name))
            self.summary_table.setItem(row, 1, QTableWidgetItem(str(total)))
            self.summary_table.setItem(row, 2, QTableWidgetItem(str(present)))
            self.summary_table.setItem(row, 3, QTableWidgetItem(str(absent)))
            self.summary_table.setItem(row, 4, QTableWidgetItem(f"{int(percent)}%"))

        # 2. Xử lý Bảng Chi Tiết
        self.detail_table.setRowCount(0)
        cursor.execute("""
            SELECT date, class_name, name, content 
            FROM progress 
            WHERE date BETWEEN ? AND ? 
            ORDER BY date DESC, class_name ASC
        """, (d1, d2))
        
        details = cursor.fetchall()
        for r_date, c_name, s_name, content in details:
            # Tách năm và tháng từ ngày của dòng hiện tại
            year = int(r_date[:4])
            month = int(r_date[5:7])
            current_month_str = r_date[:7] # Dạng "YYYY-MM"

            # A. Tính tổng số buổi có trong tháng (Các ngày T7, CN)
            total_weekends = self.count_weekends_in_month(year, month)
            
            # B. Tính số buổi học sinh này ĐÃ ĐI HỌC trong tháng đó
            cursor.execute("""
                SELECT COUNT(*) FROM progress 
                WHERE name = ? AND status = 'Đi học' AND date LIKE ?
            """, (s_name, f"{current_month_str}%"))
            attended = cursor.fetchone()[0] or 0
            
            attendance_ratio = f"{attended}/{total_weekends}"

            row = self.detail_table.rowCount()
            self.detail_table.insertRow(row)
            self.detail_table.setItem(row, 0, QTableWidgetItem(r_date))
            self.detail_table.setItem(row, 1, QTableWidgetItem(c_name))
            self.detail_table.setItem(row, 2, QTableWidgetItem(s_name))
            
            # Cột Chuyên cần (Ví dụ: 7/8)
            ratio_item = QTableWidgetItem(attendance_ratio)
            ratio_item.setTextAlignment(Qt.AlignCenter)
            
            # Đổi màu chữ: Nếu nghỉ quá 2 buổi thì hiện màu đỏ cảnh báo
            if total_weekends - attended >= 2:
                ratio_item.setForeground(QColor("#dc3545")) # Đỏ
            else:
                ratio_item.setForeground(QColor("#28a745")) # Xanh lá
                
            self.detail_table.setItem(row, 3, ratio_item)
            self.detail_table.setItem(row, 4, QTableWidgetItem(str(content)))

    def export_to_word(self):
        path, _ = QFileDialog.getSaveFileName(self, "Lưu báo cáo", f"Bao_cao_hoc_tap_{QDate.currentDate().toString('ddMMyy')}.docx", "Word Files (*.docx)")
        if not path:
            return

        try:
            doc = Document()
            doc.add_heading('BÁO CÁO TÌNH HÌNH HỌC TẬP', 0)
            
            # Thông tin thời gian
            time_str = f"Từ ngày: {self.start_date.date().toString('dd/MM/yyyy')} đến ngày: {self.end_date.date().toString('dd/MM/yyyy')}"
            doc.add_paragraph(time_str)

            # 1. Bảng tóm tắt
            doc.add_heading('1. Thống kê chuyên cần theo lớp', level=1)
            table1 = doc.add_table(rows=1, cols=5)
            table1.style = 'Table Grid'
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = 'Tên lớp'
            hdr_cells[1].text = 'Sĩ số'
            hdr_cells[2].text = 'Đi học'
            hdr_cells[3].text = 'Nghỉ học'
            hdr_cells[4].text = 'Tỉ lệ (%)'

            for r in range(self.summary_table.rowCount()):
                row_cells = table1.add_row().cells
                for c in range(5):
                    row_cells[c].text = self.summary_table.item(r, c).text()

            doc.add_paragraph("\n")

            # 2. Bảng chi tiết
            doc.add_heading('2. Chi tiết nội dung và nhận xét', level=1)
            table2 = doc.add_table(rows=1, cols=4)
            table2.style = 'Table Grid'
            hdr_cells2 = table2.rows[0].cells
            hdr_cells2[0].text = 'Ngày'
            hdr_cells2[1].text = 'Lớp'
            hdr_cells2[2].text = 'Học viên'
            hdr_cells2[3].text = 'Nhận xét'

            for r in range(self.detail_table.rowCount()):
                row_cells = table2.add_row().cells
                for c in range(4):
                    row_cells[c].text = self.detail_table.item(r, c).text()

            doc.save(path)
            QMessageBox.information(self, "Thành công", f"Đã xuất báo cáo tại:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Không thể xuất file: {str(e)}")

class StudentManager(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Sổ tay Python 2026 - Quản lý tiến độ")
        self.resize(1150, 800)
        self.center_window()
        self.setStyleSheet("""
            QMainWindow, QWidget { background-color: #ffffff; color: #333; }
            QTableWidget { gridline-color: #eee; border: 1px solid #ddd; }
            QHeaderView::section { background-color: #fcfcfc; border: 1px solid #ddd; font-weight: bold; }
            QPushButton { border-radius: 4px; font-weight: bold; }
        """)
        
        self.conn = sqlite3.connect('hoc_tap.db')
        self.init_db()
        self.setup_ui()
        self.load_data(is_reset=True)

    def init_db(self):
        cursor = self.conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS progress (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            date TEXT, name TEXT, class_name TEXT, 
                            status TEXT, content TEXT, is_highlighted INTEGER DEFAULT 0)''')
        self.conn.commit()

    def setup_ui(self):
        main_layout = QVBoxLayout()
        
        filter_layout = QHBoxLayout()
        
        self.search_name = QLineEdit(); self.search_name.setPlaceholderText("🔍 Tên học sinh...")
        
        self.filter_class = QComboBox()
        self.filter_class.addItems(["Tất cả lớp", "Sáng T7", "Chiều T7", "Sáng CN", "Chiều CN"])
        
        self.check_date = QComboBox()
        self.check_date.addItems(["Tất cả thời gian", "Theo ngày"])
        self.check_date.setFixedWidth(100)
        
        self.search_date = QDateEdit()
        self.search_date.setCalendarPopup(True)
        self.search_date.setDisplayFormat("yyyy-MM-dd")
        self.search_date.setDate(QDate.currentDate())
        self.search_date.setEnabled(False) 
        self.check_date.currentIndexChanged.connect(lambda i: self.search_date.setEnabled(i == 1))

        self.btn_filter = QPushButton("Lọc"); self.btn_filter.setFixedWidth(80)
        self.btn_filter.setStyleSheet("background-color: #f8f9fa; border: 1px solid #ccc; padding: 5px;")
        self.btn_filter.clicked.connect(lambda: self.load_data(is_reset=False))

        filter_layout.addWidget(QLabel("Lớp:")); filter_layout.addWidget(self.filter_class)
        filter_layout.addWidget(QLabel("Học sinh:")); filter_layout.addWidget(self.search_name)
        filter_layout.addWidget(self.check_date); filter_layout.addWidget(self.search_date)
        filter_layout.addWidget(self.btn_filter)
        main_layout.addLayout(filter_layout)

        # --- THANH CÔNG CỤ ---
        toolbar = QHBoxLayout()
        self.btn_att = QPushButton("✓ ĐIỂM DANH"); self.btn_att.setStyleSheet("background-color: #28a745; color: white;")
        self.btn_add = QPushButton("+ Thêm lẻ"); self.btn_add.setStyleSheet("background-color: #007bff; color: white;")
        self.btn_edit = QPushButton("✎ Viết nhận xét"); self.btn_edit.setStyleSheet("background-color: #ffc107; color: #222;")
        self.btn_del = QPushButton("🗑 Xóa"); self.btn_del.setStyleSheet("background-color: #dc3545; color: white;")
        self.btn_stats = QPushButton("📊 THỐNG KÊ");self.btn_stats.setFixedSize(140, 35); self.btn_stats.setStyleSheet("background-color: #17a2b8; color: white;")
        
        for btn in [self.btn_att, self.btn_add, self.btn_edit, self.btn_del]:
            btn.setFixedSize(140, 35)
            toolbar.addWidget(btn)
        
        self.btn_att.clicked.connect(self.open_attendance)
        self.btn_add.clicked.connect(self.add_entry)
        self.btn_edit.clicked.connect(self.edit_entry)
        self.btn_del.clicked.connect(self.delete_entry)
        toolbar.addStretch()
        main_layout.addLayout(toolbar)
        self.btn_stats.clicked.connect(self.open_statistics)
        toolbar.addWidget(self.btn_stats)

        # --- BẢNG DỮ LIỆU ---
        self.table = QTableWidget(); self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["ID", "Ngày", "Học sinh", "Lớp", "Trạng thái", "Nội dung bài học (Cập nhật cuối buổi)"])
        self.table.hideColumn(0); self.table.horizontalHeader().setSectionResizeMode(5, QHeaderView.Stretch)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setAlternatingRowColors(True)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        
        main_layout.addWidget(self.table)
        container = QWidget(); container.setLayout(main_layout); self.setCentralWidget(container)

    def center_window(self):
        qr = self.frameGeometry()
        cp = QScreen.availableGeometry(QApplication.primaryScreen()).center()
        qr.moveCenter(cp); self.move(qr.topLeft())
        
    def load_data(self, is_reset=False):
        cursor = self.conn.cursor()
        query = "SELECT * FROM progress WHERE name LIKE ?"
        params = [f"%{self.search_name.text()}%"]
        
        if self.filter_class.currentText() != "Tất cả lớp":
            query += " AND class_name = ?"
            params.append(self.filter_class.currentText())
            
        if self.check_date.currentIndex() == 1:
            query += " AND date = ?"
            params.append(self.search_date.date().toString("yyyy-MM-dd"))

        query += " ORDER BY date DESC, id DESC"
        cursor.execute(query, params)
        rows = cursor.fetchall()
        
        self.table.setRowCount(0)
        for row_data in rows:
            r = self.table.rowCount()
            self.table.insertRow(r)
            for c, val in enumerate(row_data[:-1]):
                item = QTableWidgetItem(str(val))
                if row_data[5] == "(Chưa có nhận xét cuối buổi)":
                    item.setForeground(QColor("#d9534f"))
                    if c == 5:
                        font = item.font()
                        font.setItalic(True)
                        item.setFont(font)
                if row_data[6] == 1: item.setBackground(QColor("#fff3cd"))
                elif row_data[6] == 2: item.setBackground(QColor("#d4edda"))
                elif row_data[6] == 3: item.setBackground(QColor("#f8d7da"))
                
                if c == 5 and val == "(Chưa có nhận xét cuối buổi)":
                    item.setForeground(QColor("#d9534f"))
                
                self.table.setItem(r, c, item)

    def open_attendance(self):
        dialog = AttendanceDialog(self, self.conn)
        if dialog.exec():
            selected = dialog.get_selected_data()
            if not selected:
                return
                
            cursor = self.conn.cursor()
            
            for name, cls, chosen_date in selected:
                cursor.execute("SELECT id FROM progress WHERE name = ? AND date = ?", (name, chosen_date))
                
                if cursor.fetchone() is None:
                    cursor.execute(
                        "INSERT INTO progress (date, name, class_name, status, content, is_highlighted) VALUES (?,?,?,?,?,?)",
                        (chosen_date, name, cls, "Đi học", "(Chưa có nhận xét cuối buổi)", 0)
                    )
            
            self.conn.commit()
            self.load_data()
            
            QMessageBox.information(self, "Thành công", f"Đã điểm danh cho {len(selected)} học sinh vào ngày {chosen_date}.")

    def add_entry(self):
        cursor = self.conn.cursor(); cursor.execute("SELECT DISTINCT name FROM progress")
        names = [r[0] for r in cursor.fetchall()]
        dialog = EntryDialog(self, student_list=names, db_conn=self.conn)
        if dialog.exec():
            self.conn.cursor().execute("INSERT INTO progress (date, name, class_name, status, content, is_highlighted) VALUES (?,?,?,?,?,?)", dialog.get_data())
            self.conn.commit(); self.load_data()

    def edit_entry(self):
        curr = self.table.currentRow()
        if curr < 0: return
        row_id = self.table.item(curr, 0).text()
        cursor = self.conn.cursor(); cursor.execute("SELECT * FROM progress WHERE id=?", (row_id,))
        row_data = list(cursor.fetchone())
        
        if row_data[5] == "(Chưa có nhận xét cuối buổi)": row_data[5] = ""
        
        cursor.execute("SELECT DISTINCT name FROM progress")
        names = [r[0] for r in cursor.fetchall()]
        
        dialog = EntryDialog(self, row_data, names, self.conn)
        if dialog.exec():
            self.conn.cursor().execute("UPDATE progress SET date=?, name=?, class_name=?, status=?, content=?, is_highlighted=? WHERE id=?", (*dialog.get_data(), row_id))
            self.conn.commit(); self.load_data()

    def delete_entry(self):
        selected_items = self.table.selectionModel().selectedRows()
        
        if not selected_items:
            msg_warn = QMessageBox(self)
            msg_warn.setWindowTitle("Lưu ý")
            msg_warn.setText("Vui lòng chọn ít nhất một dòng để xóa!")
            msg_warn.setIcon(QMessageBox.Warning)
            msg_warn.setStyleSheet("QMessageBox { background-color: white; } QPushButton { width: 80px; height: 30px; }")
            msg_warn.exec()
            return

        count = len(selected_items)
        
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Xác nhận xóa")
        msg_box.setText(f"<h3>Bạn có chắc chắn muốn xóa {count} bản ghi đã chọn?</h3>")
        msg_box.setInformativeText("Hành động này không thể hoàn tác.")
        msg_box.setIcon(QMessageBox.Question)
        msg_box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msg_box.setDefaultButton(QMessageBox.No)
        
        msg_box.setStyleSheet("""
            QMessageBox { background-color: #ffffff; }
            QLabel { color: #333; font-size: 14px; }
            QPushButton { 
                border-radius: 4px; padding: 6px 20px; font-weight: bold; min-width: 70px; 
            }
            QPushButton[text="&Yes"] { background-color: #dc3545; color: white; }
            QPushButton[text="&No"] { background-color: #f8f9fa; border: 1px solid #ccc; }
        """)

        if msg_box.exec() == QMessageBox.Yes:
            cursor = self.conn.cursor()
            try:
                for index in selected_items:
                    row_id = self.table.item(index.row(), 0).text()
                    cursor.execute("DELETE FROM progress WHERE id=?", (row_id,))
                
                self.conn.commit()
                self.load_data()
                
            except Exception as e:
                self.conn.rollback()
                QMessageBox.critical(self, "Lỗi", f"Không thể xóa dữ liệu: {e}")
                
    def open_statistics(self):
        dialog = StatisticsDialog(self, self.conn)
        dialog.exec()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = StudentManager(); window.show()
    sys.exit(app.exec())